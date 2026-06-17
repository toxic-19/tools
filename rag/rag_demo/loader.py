"""
文档加载模块
=============
支持 .txt / .md / .pdf / .docx 文件的解析与加载。
每个文档被解析为标准 Chunk 格式，携带来源元信息用于引用溯源。
"""

import os
import hashlib
from dataclasses import dataclass, field
from typing import List, Optional
from pathlib import Path


@dataclass
class Document:
    """文档标准表示"""
    content: str                           # 文档全文文本
    source: str                            # 文件路径
    filename: str                          # 文件名
    file_type: str                         # 文件类型
    metadata: dict = field(default_factory=dict)  # 额外元信息（页数等）


def _file_hash(filepath: str) -> str:
    """计算文件 MD5，用于去重"""
    h = hashlib.md5()
    with open(filepath, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()


# ---- 各格式解析器 ----

def _load_txt(filepath: str) -> Document:
    with open(filepath, "r", encoding="utf-8", errors="replace") as f:
        content = f.read()
    return Document(
        content=content,
        source=filepath,
        filename=os.path.basename(filepath),
        file_type="txt",
    )


def _load_md(filepath: str) -> Document:
    return _load_txt(filepath)  # Markdown 按纯文本处理即可


def _load_pdf(filepath: str) -> Document:
    """
    PDF 解析 —— 使用 PyMuPDF (fitz)。
    逐页提取文本，并在 metadata 中记录每页文本用于溯源。
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError("请安装 PyMuPDF: pip install PyMuPDF")

    doc = fitz.open(filepath)
    pages_text = []
    for i, page in enumerate(doc):
        text = page.get_text("text").strip()
        if text:
            pages_text.append((i + 1, text))  # (页码, 文本)
    doc.close()

    full_text = "\n\n".join(t for _, t in pages_text)
    page_map = {pg: txt for pg, txt in pages_text}

    return Document(
        content=full_text,
        source=filepath,
        filename=os.path.basename(filepath),
        file_type="pdf",
        metadata={"page_map": page_map, "total_pages": len(pages_text)},
    )


def _load_docx(filepath: str) -> Document:
    """
    DOCX 解析 —— 使用 python-docx。
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ImportError("请安装 python-docx: pip install python-docx")

    doc = DocxDocument(filepath)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    content = "\n\n".join(paragraphs)

    return Document(
        content=content,
        source=filepath,
        filename=os.path.basename(filepath),
        file_type="docx",
    )


# ---- 统一加载入口 ----

_LOADERS = {
    ".txt": _load_txt,
    ".md": _load_md,
    ".pdf": _load_pdf,
    ".docx": _load_docx,
}


def load_document(filepath: str) -> Document:
    """加载单个文档，根据扩展名自动选择解析器。"""
    filepath = str(Path(filepath).resolve())
    ext = os.path.splitext(filepath)[1].lower()
    if ext not in _LOADERS:
        raise ValueError(f"不支持的文件格式: {ext}，当前支持: {list(_LOADERS.keys())}")
    return _LOADERS[ext](filepath)


def load_directory(dirpath: str) -> List[Document]:
    """递归加载目录下所有支持的文档。"""
    docs = []
    for root, _, files in os.walk(dirpath):
        for fname in sorted(files):
            ext = os.path.splitext(fname)[1].lower()
            if ext in _LOADERS:
                fpath = os.path.join(root, fname)
                try:
                    docs.append(load_document(fpath))
                    print(f"  [OK] 已加载: {fname}")
                except Exception as e:
                    print(f"  [FAIL] 加载失败 {fname}: {e}")
    return docs
